"""Contract generator using the docx template with placeholder replacement."""

import base64
import copy
import io
import os
import logging
import re
from datetime import datetime, date
from dateutil.relativedelta import relativedelta
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, Cm

logger = logging.getLogger(__name__)

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "templates", "contract_template.docx")


def _replace_in_paragraph(paragraph, replacements: dict[str, str]):
    """Replace placeholders in a paragraph while preserving formatting.

    Handles placeholders that span multiple runs by joining run texts,
    performing replacement, and re-splitting into the original runs.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if not any(key in full_text for key in replacements):
        return

    # Apply all replacements to the full text
    new_text = full_text
    for key, value in replacements.items():
        new_text = new_text.replace(key, value)

    if new_text == full_text:
        return

    # Re-distribute the new text across the existing runs, keeping formatting
    # Put all text in the first run, clear the rest
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def _replace_in_table(table, replacements: dict[str, str]):
    """Replace placeholders in all cells of a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, replacements)


def _ordinal_suffix(day: int) -> str:
    if 11 <= day <= 13:
        return "th"
    return {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")


def _decode_signature_image(data_url: str) -> Optional[io.BytesIO]:
    """Decode a base64 data-url signature image into a BytesIO stream."""
    if not data_url or not data_url.startswith("data:image"):
        return None
    try:
        # data:image/png;base64,iVBOR...
        _, b64_data = data_url.split(",", 1)
        img_bytes = base64.b64decode(b64_data)
        return io.BytesIO(img_bytes)
    except Exception:
        return None


def generate_contract_docx(
    *,
    applicant_name: str,
    applicant_address: str,
    national_id: str = "",
    reference_number: str = "",
    product_name: str = "Hire Purchase",
    items: Optional[list[dict]] = None,
    amount: float = 0,
    term_months: int = 12,
    monthly_payment: float = 0,
    total_financed: float = 0,
    downpayment: float = 0,
    interest_and_fees: float = 0,
    interest_rate: Optional[float] = None,
    signed_at: Optional[datetime] = None,
    signature_name: str = "",
    signature_data_url: str = "",
    contact_details: str = "",
) -> io.BytesIO:
    """Generate a contract from the docx template with all placeholders filled."""
    try:
        if not os.path.exists(TEMPLATE_PATH):
            raise FileNotFoundError(f"Contract template not found at {TEMPLATE_PATH}")

        doc = Document(TEMPLATE_PATH)

        # ── Compute date values ───────────────────────────
        now = signed_at or datetime.now()
        day = now.day
        day_suffix = _ordinal_suffix(day)
        month_name = now.strftime("%B")
        year = str(now.year)
        date_str = now.strftime("%d/%m/%Y")

        expiry_date = now + relativedelta(months=term_months)
        expiry_str = expiry_date.strftime("%d/%m/%Y")

        # Repayment day — typically the same day of month as start
        repayment_day = f"{day}{day_suffix}"

        total_repayment = monthly_payment * term_months if monthly_payment > 0 else total_financed
        if interest_and_fees <= 0:
            interest_and_fees = total_repayment - (amount - downpayment) if total_repayment > (amount - downpayment) else 0

        # Hire Purchase Price = total of all payments (principal - downpayment + credit charges)
        hire_purchase_price = total_repayment if total_repayment > 0 else total_financed

        # ── Build replacements dict ───────────────────────
        replacements: dict[str, str] = {
            "{NAME}": applicant_name or "________________________",
            "{ADDRESS}": applicant_address or "________________________",
            "{DAY}": f"{day}{day_suffix}",
            "{MONTH}": month_name,
            "{YEAR}": year,
            "{DATE}": date_str,
            "{DATE_NAME}": f"{date_str} — {applicant_name}" if applicant_name else date_str,
            "{TOTAL LOAN AMOUNT}": f"TTD {hire_purchase_price:,.2f}",
            "{PRINCIPAL}": f"TTD {amount:,.2f}",
            "{DOWNPAYMENT}": f"TTD {downpayment:,.2f}",
            "{INTEREST AND FEES}": f"TTD {interest_and_fees:,.2f}",
            "{TOTAL}": f"TTD {hire_purchase_price:,.2f}",
            "{TENURE}": str(term_months),
            "{INSTALMENT AMOUNT}": f"TTD {monthly_payment:,.2f}",
            "{INSTALMENT}": f"TTD {monthly_payment:,.2f}",
            "{DATE + TENURE}": expiry_str,
            "{REPAYMENT DAY}": repayment_day,
            "{CONTACT_DETAILS}": contact_details or "________________________",
        }

        # Handle items in Schedule 1 table
        # Items table is Table index 2
        # Template rows: 0 = header, 1 = first item, 2 = second item placeholder
        if items and len(items) > 0 and len(doc.tables) > 2:
            items_table = doc.tables[2]

            # Fill item rows
            for idx, item in enumerate(items):
                item_name = item.get("description") or item.get("category_name") or f"Item {idx + 1}"
                item_qty = str(item.get("quantity", 1))
                item_price = f"TTD {float(item.get('price', 0)):,.2f}"
                order_num = str(idx + 1)

                if idx < 2 and (idx + 1) < len(items_table.rows):
                    # Replace in existing template rows
                    row = items_table.rows[idx + 1]
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            text = para.text
                            text = text.replace("{ORDER NUMBER}", order_num)
                            text = text.replace("{ITEM 1}", item_name)
                            text = text.replace("{ITEM …}", item_name)
                            text = text.replace("{QUANTITY}", item_qty)
                            if para.runs:
                                para.runs[0].text = text
                                for run in para.runs[1:]:
                                    run.text = ""
                    # Set the value cell (last column)
                    row.cells[-1].text = item_price
                else:
                    # Add new row for additional items
                    new_row = items_table.add_row()
                    new_row.cells[0].text = order_num
                    new_row.cells[1].text = ""
                    new_row.cells[2].text = item_name
                    new_row.cells[3].text = item_qty
                    new_row.cells[4].text = item_price

            # If only 1 item, remove the second placeholder row (row index 2)
            if len(items) == 1 and len(items_table.rows) > 2:
                placeholder_row = items_table.rows[2]
                tbl_element = items_table._tbl
                tbl_element.remove(placeholder_row._tr)
        else:
            # No items — remove the second placeholder row and clean up the first
            if len(doc.tables) > 2:
                items_table = doc.tables[2]
                # Remove second placeholder row
                if len(items_table.rows) > 2:
                    placeholder_row = items_table.rows[2]
                    items_table._tbl.remove(placeholder_row._tr)
            replacements["{ORDER NUMBER}"] = ""
            replacements["{ITEM 1}"] = product_name
            replacements["{ITEM …}"] = ""
            replacements["{QUANTITY}"] = "1"

        # ── Apply replacements to all paragraphs ──────────
        for paragraph in doc.paragraphs:
            _replace_in_paragraph(paragraph, replacements)

        # ── Apply replacements to all tables ──────────────
        for table in doc.tables:
            _replace_in_table(table, replacements)

        # ── Embed signature image if available ────────────
        # The signature goes into the "Hirer Signature" cell in the signing table (Table 1)
        sig_image = _decode_signature_image(signature_data_url)
        if sig_image and len(doc.tables) > 1:
            signing_table = doc.tables[1]
            # Row 1 is "Signature: ..." row; Hirer column is index 1
            if len(signing_table.rows) > 1 and len(signing_table.columns) > 1:
                sig_cell = signing_table.rows[1].cells[1]
                sig_cell.text = ""  # clear placeholder
                sig_para = sig_cell.paragraphs[0] if sig_cell.paragraphs else sig_cell.add_paragraph()
                sig_para.add_run("Signature: ")
                sig_para.add_run().add_picture(sig_image, width=Cm(5))

        # ── Save to buffer ────────────────────────────────
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        logger.exception(f"Error in generate_contract_docx: {e}")
        raise
