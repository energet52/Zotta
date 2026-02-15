"""Convert DOCX BytesIO to PDF BytesIO.

Strategy:
1. Try LibreOffice headless (best fidelity) if available.
2. Fall back to pure-Python conversion via reportlab (always available).
"""

import io
import os
import re
import shutil
import subprocess
import tempfile
import logging
import base64

logger = logging.getLogger(__name__)


# ── LibreOffice approach ──────────────────────────────────────────────────

def _find_libreoffice() -> str | None:
    """Find the LibreOffice binary on the system."""
    for name in ("libreoffice", "soffice", "libreoffice7.6"):
        path = shutil.which(name)
        if path:
            return path
    for candidate in (
        "/usr/bin/libreoffice",
        "/usr/bin/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ):
        if os.path.isfile(candidate):
            return candidate
    return None


def _convert_with_libreoffice(docx_buffer: io.BytesIO) -> io.BytesIO:
    """Convert DOCX to PDF using LibreOffice headless."""
    lo_bin = _find_libreoffice()
    if not lo_bin:
        raise RuntimeError("LibreOffice not found")

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "document.docx")
        pdf_path = os.path.join(tmpdir, "document.pdf")

        with open(docx_path, "wb") as f:
            f.write(docx_buffer.getvalue())

        env = os.environ.copy()
        env["HOME"] = tmpdir

        result = subprocess.run(
            [lo_bin, "--headless", "--norestore", "--convert-to", "pdf",
             "--outdir", tmpdir, docx_path],
            capture_output=True, text=True, timeout=60, env=env,
        )

        if result.returncode != 0 or not os.path.exists(pdf_path):
            raise RuntimeError(f"LibreOffice failed (rc={result.returncode}): {result.stderr}")

        with open(pdf_path, "rb") as f:
            buf = io.BytesIO(f.read())
            buf.seek(0)
            return buf


# ── Pure-Python (reportlab) approach ──────────────────────────────────────

def _convert_with_reportlab(docx_buffer: io.BytesIO) -> io.BytesIO:
    """Convert DOCX to PDF by extracting text+tables via python-docx and
    rendering with reportlab.  Not pixel-perfect but fully functional."""
    from docx import Document
    from docx.table import Table
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table as RLTable,
        TableStyle, Image, PageBreak,
    )
    from reportlab.lib import colors

    doc = Document(docx_buffer)
    styles = getSampleStyleSheet()

    # Custom styles
    styles.add(ParagraphStyle(
        "ContractTitle", parent=styles["Heading1"],
        fontSize=14, alignment=TA_CENTER, spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        "ContractSubtitle", parent=styles["Normal"],
        fontSize=8, alignment=TA_CENTER, textColor=colors.grey,
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        "ContractHeading", parent=styles["Heading2"],
        fontSize=11, spaceBefore=10, spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        "ContractBody", parent=styles["Normal"],
        fontSize=9, leading=12, alignment=TA_JUSTIFY, spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        "ContractBodyBold", parent=styles["Normal"],
        fontSize=9, leading=12, alignment=TA_JUSTIFY, spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        "ContractCenter", parent=styles["Normal"],
        fontSize=9, alignment=TA_CENTER, spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        "ContractSmall", parent=styles["Normal"],
        fontSize=7, leading=9, alignment=TA_JUSTIFY, spaceAfter=2,
    ))

    story: list = []

    def _escape(text: str) -> str:
        """Escape XML special chars for reportlab Paragraph."""
        return (text
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace('"', "&quot;")
                .replace("'", "&#39;"))

    def _guess_style(para) -> str:
        """Pick a reportlab style based on the docx paragraph properties."""
        text = para.text.strip()
        if not text:
            return "__skip__"

        # Check paragraph style name
        sname = (para.style.name or "").lower() if para.style else ""

        # Check if bold (all runs bold or style is heading)
        all_bold = all(r.bold for r in para.runs if r.text.strip()) if para.runs else False

        if "heading" in sname or "title" in sname:
            if "1" in sname or "title" in sname:
                return "ContractTitle"
            return "ContractHeading"

        # Detect centered
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            if all_bold:
                return "ContractTitle"
            return "ContractCenter"

        # Numbered clauses or bold lines
        if all_bold and len(text) < 120:
            return "ContractHeading"

        # Short uppercase => subtitle
        if text.isupper() and len(text) < 100:
            return "ContractCenter"

        return "ContractBody"

    def _para_to_markup(para) -> str:
        """Convert a docx paragraph to reportlab Paragraph markup with bold/italic."""
        parts: list[str] = []
        for run in para.runs:
            t = _escape(run.text)
            if not t:
                continue
            if run.bold and run.italic:
                t = f"<b><i>{t}</i></b>"
            elif run.bold:
                t = f"<b>{t}</b>"
            elif run.italic:
                t = f"<i>{t}</i>"
            parts.append(t)
        return "".join(parts) if parts else _escape(para.text)

    def _add_table(table: Table):
        """Convert a docx table to a reportlab Table."""
        data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(Paragraph(_escape(cell_text), styles["ContractSmall"]))
            data.append(row_data)

        if not data:
            return

        # Calculate column widths (equal distribution)
        num_cols = max(len(r) for r in data) if data else 1
        available = 170 * mm
        col_widths = [available / num_cols] * num_cols

        rl_table = RLTable(data, colWidths=col_widths, repeatRows=1)
        rl_table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.93, 0.93, 0.93)),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(rl_table)
        story.append(Spacer(1, 4 * mm))

    # ── Walk through document body (paragraphs + tables in order) ──
    # python-docx body elements include both paragraphs and tables
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Find the matching paragraph object
            for para in doc.paragraphs:
                if para._element is element:
                    style_name = _guess_style(para)
                    if style_name == "__skip__":
                        story.append(Spacer(1, 2 * mm))
                    else:
                        markup = _para_to_markup(para)
                        if markup.strip():
                            story.append(Paragraph(markup, styles[style_name]))
                    break

        elif tag == "tbl":
            for table in doc.tables:
                if table._element is element:
                    _add_table(table)
                    break

    # ── Check for signature images in the DOCX ──
    # Look for images in relationships
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_data = rel.target_part.blob
                img_buf = io.BytesIO(img_data)
                img = Image(img_buf, width=50 * mm, height=15 * mm)
                story.append(Spacer(1, 3 * mm))
                story.append(Paragraph("<b>Signature:</b>", styles["ContractBody"]))
                story.append(img)
            except Exception:
                pass  # skip if image fails

    # ── Build PDF ──
    pdf_buffer = io.BytesIO()
    pdf_doc = SimpleDocTemplate(
        pdf_buffer,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
        title="Hire Purchase Agreement",
    )

    if not story:
        story.append(Paragraph("Contract document", styles["ContractBody"]))

    pdf_doc.build(story)
    pdf_buffer.seek(0)
    return pdf_buffer


# ── Public API ────────────────────────────────────────────────────────────

def convert_docx_to_pdf(docx_buffer: io.BytesIO) -> io.BytesIO:
    """Convert a DOCX file (as BytesIO) to PDF.

    Tries LibreOffice first (best quality), falls back to reportlab
    (pure Python, always works).
    """
    # Try LibreOffice first
    try:
        return _convert_with_libreoffice(docx_buffer)
    except Exception as e:
        logger.info("LibreOffice not available (%s), using reportlab fallback", e)

    # Fall back to reportlab
    docx_buffer.seek(0)
    return _convert_with_reportlab(docx_buffer)
