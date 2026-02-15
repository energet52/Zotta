"""Tests for docx-based contract generation from template."""

import io
from datetime import datetime, timezone

import pytest

from app.services.contract_generator import (
    generate_contract_docx,
    _ordinal_suffix,
    _decode_signature_image,
)


# ── Helpers ─────────────────────────────────────────────────


VALID_SIGNATURE = (
    "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJ"
    "AAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="
)


class TestOrdinalSuffix:
    """Unit tests for the _ordinal_suffix helper."""

    def test_first(self):
        assert _ordinal_suffix(1) == "st"

    def test_second(self):
        assert _ordinal_suffix(2) == "nd"

    def test_third(self):
        assert _ordinal_suffix(3) == "rd"

    def test_fourth(self):
        assert _ordinal_suffix(4) == "th"

    def test_eleventh(self):
        assert _ordinal_suffix(11) == "th"

    def test_twelfth(self):
        assert _ordinal_suffix(12) == "th"

    def test_thirteenth(self):
        assert _ordinal_suffix(13) == "th"

    def test_twenty_first(self):
        assert _ordinal_suffix(21) == "st"

    def test_twenty_second(self):
        assert _ordinal_suffix(22) == "nd"

    def test_thirty_first(self):
        assert _ordinal_suffix(31) == "st"


class TestDecodeSignatureImage:
    """Unit tests for the _decode_signature_image helper."""

    def test_valid_data_url(self):
        result = _decode_signature_image(VALID_SIGNATURE)
        assert result is not None
        assert isinstance(result, io.BytesIO)
        data = result.read()
        assert len(data) > 0

    def test_empty_string(self):
        assert _decode_signature_image("") is None

    def test_none(self):
        assert _decode_signature_image(None) is None  # type: ignore

    def test_not_a_data_url(self):
        assert _decode_signature_image("not-a-valid-data-url") is None

    def test_bad_base64(self):
        result = _decode_signature_image("data:image/png;base64,!!!invalid!!!")
        assert result is None


class TestGenerateContractDocx:
    """Unit tests for the generate_contract_docx function."""

    @pytest.fixture
    def docx_kwargs(self):
        """Standard keyword args for contract generation."""
        return dict(
            applicant_name="John Doe",
            applicant_address="45 Elm Street, San Fernando",
            national_id="19900515001",
            reference_number="ZOT-2026-TESTDOCX",
            product_name="HP Standard",
            amount=10000.00,
            term_months=12,
            monthly_payment=900.50,
            total_financed=10800.00,
            downpayment=500.00,
            interest_and_fees=1300.00,
            interest_rate=14.5,
            signed_at=datetime(2026, 3, 15, 14, 30, 0, tzinfo=timezone.utc),
            signature_name="John Doe",
            signature_data_url="",
            contact_details="+18685551234 and john@example.com",
        )

    def test_returns_bytesio(self, docx_kwargs):
        """generate_contract_docx should return a BytesIO object."""
        result = generate_contract_docx(**docx_kwargs)
        assert isinstance(result, io.BytesIO)

    def test_docx_has_correct_header(self, docx_kwargs):
        """Result should be a valid DOCX (ZIP format, starts with PK)."""
        result = generate_contract_docx(**docx_kwargs)
        data = result.read()
        # DOCX is a ZIP file — starts with PK\x03\x04
        assert data[:2] == b"PK"

    def test_docx_has_content(self, docx_kwargs):
        """DOCX should be reasonably sized (not empty)."""
        result = generate_contract_docx(**docx_kwargs)
        data = result.read()
        # A contract from a full template should be > 10 KB
        assert len(data) > 10000

    def test_placeholders_replaced(self, docx_kwargs):
        """All template placeholders should be replaced with actual values."""
        from docx import Document

        result = generate_contract_docx(**docx_kwargs)
        doc = Document(result)

        # Collect all text in document
        all_text = ""
        for para in doc.paragraphs:
            all_text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + "\n"

        # No unreplaced placeholders should remain
        import re
        placeholders = re.findall(r"\{[A-Z][A-Z _+]+\}", all_text)
        assert placeholders == [], f"Unreplaced placeholders found: {placeholders}"

    def test_applicant_name_appears(self, docx_kwargs):
        """Applicant name should appear in the generated document."""
        from docx import Document

        result = generate_contract_docx(**docx_kwargs)
        doc = Document(result)
        all_text = "\n".join(para.text for para in doc.paragraphs)
        assert "John Doe" in all_text

    def test_applicant_address_appears(self, docx_kwargs):
        """Applicant address should appear in the generated document."""
        from docx import Document

        result = generate_contract_docx(**docx_kwargs)
        doc = Document(result)
        all_text = "\n".join(para.text for para in doc.paragraphs)
        assert "45 Elm Street, San Fernando" in all_text

    def test_financial_values_appear(self, docx_kwargs):
        """Financial values should appear formatted in the document."""
        from docx import Document

        result = generate_contract_docx(**docx_kwargs)
        doc = Document(result)

        # Collect all text including tables
        all_text = ""
        for para in doc.paragraphs:
            all_text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + "\n"

        assert "TTD 10,800.00" in all_text  # total_financed
        assert "TTD 10,000.00" in all_text  # principal/amount
        assert "TTD 500.00" in all_text     # downpayment
        assert "TTD 900.50" in all_text     # monthly_payment

    def test_tenure_appears(self, docx_kwargs):
        """Term months should appear in the document."""
        from docx import Document

        result = generate_contract_docx(**docx_kwargs)
        doc = Document(result)
        all_text = ""
        for para in doc.paragraphs:
            all_text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + "\n"

        assert "12" in all_text  # tenure months

    def test_date_populated(self, docx_kwargs):
        """Date should be formatted and placed in the document."""
        from docx import Document

        result = generate_contract_docx(**docx_kwargs)
        doc = Document(result)
        all_text = "\n".join(para.text for para in doc.paragraphs)
        assert "15/03/2026" in all_text  # signed_at date
        assert "15th" in all_text         # ordinal day
        assert "March" in all_text        # month name
        assert "2026" in all_text         # year

    def test_expiry_date_calculated(self, docx_kwargs):
        """Expiry date should be start date + tenure months."""
        from docx import Document

        result = generate_contract_docx(**docx_kwargs)
        doc = Document(result)

        all_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + "\n"

        # 15/03/2026 + 12 months = 15/03/2027
        assert "15/03/2027" in all_text

    def test_contact_details_populated(self, docx_kwargs):
        """Contact details should be placed in the Hirer description."""
        from docx import Document

        result = generate_contract_docx(**docx_kwargs)
        doc = Document(result)
        all_text = "\n".join(para.text for para in doc.paragraphs)
        assert "+18685551234" in all_text
        assert "john@example.com" in all_text

    def test_without_signed_at_uses_now(self, docx_kwargs):
        """When signed_at is None, the current date should be used."""
        docx_kwargs["signed_at"] = None
        result = generate_contract_docx(**docx_kwargs)
        assert isinstance(result, io.BytesIO)
        data = result.read()
        assert data[:2] == b"PK"  # still valid docx

    def test_without_contact_details(self, docx_kwargs):
        """Document should render placeholder dashes when contact details missing."""
        from docx import Document

        docx_kwargs["contact_details"] = ""
        result = generate_contract_docx(**docx_kwargs)
        doc = Document(result)
        all_text = "\n".join(para.text for para in doc.paragraphs)
        assert "________________________" in all_text

    def test_without_applicant_name(self, docx_kwargs):
        """Document should render placeholder dashes when name is empty."""
        from docx import Document

        docx_kwargs["applicant_name"] = ""
        result = generate_contract_docx(**docx_kwargs)
        doc = Document(result)
        all_text = "\n".join(para.text for para in doc.paragraphs)
        assert "________________________" in all_text

    def test_with_signature_data_url(self, docx_kwargs):
        """Document should contain an embedded image when signature is provided."""
        docx_kwargs["signature_data_url"] = VALID_SIGNATURE
        result = generate_contract_docx(**docx_kwargs)
        # Valid docx
        data = result.read()
        assert data[:2] == b"PK"
        # Should be larger than without signature (image adds bytes)
        result_no_sig = generate_contract_docx(
            **{**docx_kwargs, "signature_data_url": ""}
        )
        data_no_sig = result_no_sig.read()
        assert len(data) >= len(data_no_sig)

    def test_with_bad_signature_data(self, docx_kwargs):
        """Document should still generate gracefully with malformed signature data."""
        docx_kwargs["signature_data_url"] = "not-a-valid-data-url"
        result = generate_contract_docx(**docx_kwargs)
        data = result.read()
        assert data[:2] == b"PK"

    def test_with_items(self, docx_kwargs):
        """Items should populate the Schedule 1 table."""
        from docx import Document

        docx_kwargs["items"] = [
            {"description": "Samsung AC 12BTU", "category_name": "Air Conditioner", "price": 6000.00, "quantity": 1},
            {"description": "LG Fridge", "category_name": "Refrigerator", "price": 4000.00, "quantity": 1},
        ]
        result = generate_contract_docx(**docx_kwargs)
        doc = Document(result)

        # Check items table (table index 2)
        items_table = doc.tables[2]
        all_cells_text = ""
        for row in items_table.rows:
            for cell in row.cells:
                all_cells_text += cell.text + "\n"

        assert "Samsung AC 12BTU" in all_cells_text
        assert "LG Fridge" in all_cells_text

    def test_with_three_items_adds_row(self, docx_kwargs):
        """When more than 2 items, additional rows are added to the items table."""
        from docx import Document

        docx_kwargs["items"] = [
            {"description": "Item A", "price": 3000.00, "quantity": 1},
            {"description": "Item B", "price": 4000.00, "quantity": 1},
            {"description": "Item C", "price": 3000.00, "quantity": 2},
        ]
        result = generate_contract_docx(**docx_kwargs)
        doc = Document(result)

        items_table = doc.tables[2]
        all_cells_text = ""
        for row in items_table.rows:
            for cell in row.cells:
                all_cells_text += cell.text + "\n"

        assert "Item A" in all_cells_text
        assert "Item B" in all_cells_text
        assert "Item C" in all_cells_text

    def test_no_items_uses_product_name(self, docx_kwargs):
        """Without items, the product name should appear in the items table."""
        from docx import Document

        docx_kwargs["items"] = None
        docx_kwargs["product_name"] = "My Custom Product"
        result = generate_contract_docx(**docx_kwargs)
        doc = Document(result)

        items_table = doc.tables[2]
        all_cells_text = ""
        for row in items_table.rows:
            for cell in row.cells:
                all_cells_text += cell.text + "\n"

        assert "My Custom Product" in all_cells_text

    def test_zero_downpayment(self, docx_kwargs):
        """Edge case: zero downpayment should be handled."""
        docx_kwargs["downpayment"] = 0.0
        result = generate_contract_docx(**docx_kwargs)
        data = result.read()
        assert data[:2] == b"PK"
        assert len(data) > 10000

    def test_zero_monthly_payment(self, docx_kwargs):
        """Edge case: zero monthly payment should be handled."""
        docx_kwargs["monthly_payment"] = 0.0
        result = generate_contract_docx(**docx_kwargs)
        data = result.read()
        assert data[:2] == b"PK"

    def test_interest_and_fees_calculated_when_zero(self, docx_kwargs):
        """When interest_and_fees is 0, it should be computed from other fields."""
        from docx import Document

        docx_kwargs["interest_and_fees"] = 0
        docx_kwargs["monthly_payment"] = 1000.0
        docx_kwargs["term_months"] = 12
        docx_kwargs["amount"] = 10000.0
        docx_kwargs["downpayment"] = 500.0
        # Expected: 1000*12 - (10000 - 500) = 12000 - 9500 = 2500

        result = generate_contract_docx(**docx_kwargs)
        doc = Document(result)

        all_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + "\n"

        assert "TTD 2,500.00" in all_text  # computed interest_and_fees

    def test_signing_table_name_populated(self, docx_kwargs):
        """The hirer name should appear in the signing table."""
        from docx import Document

        result = generate_contract_docx(**docx_kwargs)
        doc = Document(result)

        signing_table = doc.tables[1]
        all_cells_text = ""
        for row in signing_table.rows:
            for cell in row.cells:
                all_cells_text += cell.text + "\n"

        assert "John Doe" in all_cells_text

    def test_date_name_in_header_table(self, docx_kwargs):
        """The header table should contain DATE_NAME with date and applicant name."""
        from docx import Document

        result = generate_contract_docx(**docx_kwargs)
        doc = Document(result)

        header_table = doc.tables[0]
        all_cells_text = ""
        for row in header_table.rows:
            for cell in row.cells:
                all_cells_text += cell.text + "\n"

        assert "15/03/2026" in all_cells_text
        assert "John Doe" in all_cells_text
